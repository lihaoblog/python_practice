from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# 实例化word对象
doc=Document()
for i in range(3):
    #添加标题
    doc.add_paragraph('这里添加每页标题')

    #创建表格
    table=doc.add_table(3,4,'Table Grid')
    #写第0行数据
    table.rows[0].cells[0].text='1'
    table.rows[0].cells[1].text='20250603'
    table.rows[0].cells[2].text='俱乐部'
    table.rows[0].cells[3].text='【Xx字111号】'
    #写第1行数据
    table.rows[1].cells[1].text = '注意安全通知'
    # 写第2行数据
    table.rows[2].cells[1].text = '签字'

    #合并单元格，合并时候要注意不能多次合并，上一次合并的下一句要忽略这一个
    table.cell(0,0).merge(table.cell(2,0))
    table.cell(1,1).merge(table.cell(1,3))
    table.cell(2,1).merge(table.cell(2,3))

    #居中处理
    table.cell(1,1).paragraphs[0].paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER

    #设置分页
    doc.add_page_break()


doc.save(r'C:\Users\lihao\Desktop\dify\表格转word.docx')
print('success')

