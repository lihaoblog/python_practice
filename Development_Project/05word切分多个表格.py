"""
    一边是需要填写信息的合同模板，一边是记录信息的excel表格，要把这些信息全部写入到word文档中
"""
from docx import Document
import pandas as pd



# 取信息表格的每个列名
df = pd.read_excel(r'C:\Users\lihao\Desktop\dify\测试合同的填写内容.xlsx')
# print(df.columns)  #取出列名
# print(df.loc[0,'合同编号'])   #取出合同编号下的第0行的值
# '合同编号':'[合同编号]'.replace('[合同编号]','000')

# 1.读取每个word合同模板的段落，让后续来写入
for i in range(0,len(df)):    #每一行
    doc = Document(r'C:\Users\lihao\Desktop\dify\表格形式合同范本.docx')
    all_paragraph = doc.paragraphs    #读取word每个段落
    all_table = doc.tables

    for j in df.columns:        #读填空表格的每一列
        # print(df.loc[i,j])

        # 2.读出每个信息表格的数据，后续来换
        # run 代表同一段落中连续、格式相同的一段文字，比如同一字体、同一粗细、同一颜色，
        # 3.把每个段落的缺失替换成表格的数据
        for paragraph in all_paragraph:
            for run in paragraph.runs:
                run_text=run.text.replace(j,df.loc[i,j])
                run.text=run_text #这个是run_text全文，每次替换的是其中一块，所以要赋值回去继续下一块，否则永远只换第一块
        # 打印表格的数据
        all_table = doc.tables
        # 遍历每个切分后的表格
        for table in all_table:
            # 取每一行
            for row in table.rows:
                # 取每行的每个格子，也叫每个数据
                for cell in row.cells:
                    cell_text=cell.text.replace(j,df.loc[i,j])
                    # 这个是cell_text全文，每次替换的是其中一块，所以要赋值回去继续下一块，否则永远只换第一块
                    cell.text=cell_text
    doc.save(fr'C:\Users\lihao\Desktop\dify\{df.loc[i,'合同编号']}合同.docx')
    print('success')